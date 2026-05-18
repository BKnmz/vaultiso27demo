"""
VaultISO27 — backend core: constants, data helpers, export functions.
Imported by app.py (routing shell) and all page modules.
"""
__version__ = "0.4.0"
import io
import json
import platform
import re
import subprocess
import sys
import time
from datetime import datetime
from pathlib import Path

import requests
import streamlit as st
import yaml

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
BASE_DIR    = Path(__file__).parent.parent
CONFIG_PATH = BASE_DIR / "config.yaml"
ORG_PATH    = BASE_DIR / "inputs" / "organization_data.json"
OUTPUTS_DIR = BASE_DIR / "outputs"
SKILLS_DIR  = BASE_DIR / "skills"
LOG_FILE    = BASE_DIR / "logs" / "vaultiso.log"

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
CLAUSE_NAMES = {
    "4.1":   "Context of the Organization",
    "4.2":   "Interested Parties",
    "4.3":   "Scope",
    "5.1":   "Leadership Commitment",
    "5.2":   "Information Security Policy",
    "5.3":   "Roles and Responsibilities",
    "6.1":   "Risk & Opportunities Framework",
    "6.1.2": "Risk Assessment Procedure",
    "6.1.3": "Risk Treatment Plan",
    "6.2":   "Security Objectives",
    "7.1":   "Resources",
    "7.2":   "Competence",
    "7.3":   "Awareness Program",
    "7.4":   "Communication Plan",
    "7.5":   "Documented Information",
    "8.1":   "Operational Planning",
    "8.2":   "Risk Assessment (Operational)",
    "8.3":   "Risk Treatment (Operational)",
    "9.1":   "Monitoring & Measurement",
    "9.2":   "Internal Audit",
    "9.3":   "Management Review",
    "10.1":  "Nonconformity & Corrective Action",
    "10.2":  "Continual Improvement",
}

# New: kind strings for pill rendering (replaces emoji STATUS_ICON)
STATUS_KIND = {"APPROVED": "ok", "DRAFT": "warn", "REVISION": "err", "MISSING": "neutral"}

# Internal review result → (label, streamlit_type, emoji) — keep 3-tuple for export_all_to_excel compat
REVIEW_RESULT = {
    "PASS":             ("Approved by AI",    "success", "ok"),
    "CONDITIONAL PASS": ("Needs Minor Fixes", "warning", "warn"),
    "FAIL":             ("Needs Revision",    "error",   "err"),
    "UNKNOWN":          ("Not Reviewed",      "info",    "neutral"),
}

MANDATORY_ORG_FIELDS = [
    ("name",              "Company name"),
    ("industry",          "Industry / sector"),
    ("size",              "Number of employees"),
    ("scope",             "What the company does (ISMS scope)"),
    ("locations",         "Where the company operates"),
    ("legal_basis",       "Data privacy / regulatory laws that apply (e.g. GDPR, NIS2, SOC2, PCI-DSS)"),
    ("primary_processes", "Main business processes"),
    ("existing_controls", "Existing security measures (e.g. antivirus, VPN)"),
]

# ---------------------------------------------------------------------------
# ISO 27001:2022 Annex A — all 93 controls
# ---------------------------------------------------------------------------
ANNEX_A_THEMES = {
    "5": "Organizational controls",
    "6": "People controls",
    "7": "Physical controls",
    "8": "Technological controls",
}

ANNEX_A_CONTROLS: dict[str, dict] = {
    # ── Theme 5: Organizational (37) ─────────────────────────────────────
    "5.1":  {"theme": "5", "name": "Policies for information security"},
    "5.2":  {"theme": "5", "name": "Information security roles and responsibilities"},
    "5.3":  {"theme": "5", "name": "Segregation of duties"},
    "5.4":  {"theme": "5", "name": "Management responsibilities"},
    "5.5":  {"theme": "5", "name": "Contact with authorities"},
    "5.6":  {"theme": "5", "name": "Contact with special interest groups"},
    "5.7":  {"theme": "5", "name": "Threat intelligence"},
    "5.8":  {"theme": "5", "name": "Information security in project management"},
    "5.9":  {"theme": "5", "name": "Inventory of information and other associated assets"},
    "5.10": {"theme": "5", "name": "Acceptable use of information and other associated assets"},
    "5.11": {"theme": "5", "name": "Return of assets"},
    "5.12": {"theme": "5", "name": "Classification of information"},
    "5.13": {"theme": "5", "name": "Labelling of information"},
    "5.14": {"theme": "5", "name": "Information transfer"},
    "5.15": {"theme": "5", "name": "Access control"},
    "5.16": {"theme": "5", "name": "Identity management"},
    "5.17": {"theme": "5", "name": "Authentication information"},
    "5.18": {"theme": "5", "name": "Access rights"},
    "5.19": {"theme": "5", "name": "Information security in supplier relationships"},
    "5.20": {"theme": "5", "name": "Addressing information security within supplier agreements"},
    "5.21": {"theme": "5", "name": "Managing information security in the ICT supply chain"},
    "5.22": {"theme": "5", "name": "Monitoring, review and change management of supplier services"},
    "5.23": {"theme": "5", "name": "Information security for use of cloud services"},
    "5.24": {"theme": "5", "name": "Information security incident management planning and preparation"},
    "5.25": {"theme": "5", "name": "Assessment and decision on information security events"},
    "5.26": {"theme": "5", "name": "Response to information security incidents"},
    "5.27": {"theme": "5", "name": "Learning from information security incidents"},
    "5.28": {"theme": "5", "name": "Collection of evidence"},
    "5.29": {"theme": "5", "name": "Information security during disruption"},
    "5.30": {"theme": "5", "name": "ICT readiness for business continuity"},
    "5.31": {"theme": "5", "name": "Legal, statutory, regulatory and contractual requirements"},
    "5.32": {"theme": "5", "name": "Intellectual property rights"},
    "5.33": {"theme": "5", "name": "Protection of records"},
    "5.34": {"theme": "5", "name": "Privacy and protection of personally identifiable information"},
    "5.35": {"theme": "5", "name": "Independent review of information security"},
    "5.36": {"theme": "5", "name": "Compliance with policies, rules and standards for information security"},
    "5.37": {"theme": "5", "name": "Documented operating procedures"},
    # ── Theme 6: People (8) ───────────────────────────────────────────────
    "6.1":  {"theme": "6", "name": "Screening"},
    "6.2":  {"theme": "6", "name": "Terms and conditions of employment"},
    "6.3":  {"theme": "6", "name": "Information security awareness, education and training"},
    "6.4":  {"theme": "6", "name": "Disciplinary process"},
    "6.5":  {"theme": "6", "name": "Responsibilities after termination or change of employment"},
    "6.6":  {"theme": "6", "name": "Confidentiality or non-disclosure agreements"},
    "6.7":  {"theme": "6", "name": "Remote working"},
    "6.8":  {"theme": "6", "name": "Information security event reporting"},
    # ── Theme 7: Physical (14) ────────────────────────────────────────────
    "7.1":  {"theme": "7", "name": "Physical security perimeters"},
    "7.2":  {"theme": "7", "name": "Physical entry"},
    "7.3":  {"theme": "7", "name": "Securing offices, rooms and facilities"},
    "7.4":  {"theme": "7", "name": "Physical security monitoring"},
    "7.5":  {"theme": "7", "name": "Protecting against physical and environmental threats"},
    "7.6":  {"theme": "7", "name": "Working in secure areas"},
    "7.7":  {"theme": "7", "name": "Clear desk and clear screen"},
    "7.8":  {"theme": "7", "name": "Equipment siting and protection"},
    "7.9":  {"theme": "7", "name": "Security of assets off-premises"},
    "7.10": {"theme": "7", "name": "Storage media"},
    "7.11": {"theme": "7", "name": "Supporting utilities"},
    "7.12": {"theme": "7", "name": "Cabling security"},
    "7.13": {"theme": "7", "name": "Equipment maintenance"},
    "7.14": {"theme": "7", "name": "Secure disposal or re-use of equipment"},
    # ── Theme 8: Technological (34) ───────────────────────────────────────
    "8.1":  {"theme": "8", "name": "User endpoint devices"},
    "8.2":  {"theme": "8", "name": "Privileged access rights"},
    "8.3":  {"theme": "8", "name": "Information access restriction"},
    "8.4":  {"theme": "8", "name": "Access to source code"},
    "8.5":  {"theme": "8", "name": "Secure authentication"},
    "8.6":  {"theme": "8", "name": "Capacity management"},
    "8.7":  {"theme": "8", "name": "Protection against malware"},
    "8.8":  {"theme": "8", "name": "Management of technical vulnerabilities"},
    "8.9":  {"theme": "8", "name": "Configuration management"},
    "8.10": {"theme": "8", "name": "Information deletion"},
    "8.11": {"theme": "8", "name": "Data masking"},
    "8.12": {"theme": "8", "name": "Data leakage prevention"},
    "8.13": {"theme": "8", "name": "Information backup"},
    "8.14": {"theme": "8", "name": "Redundancy of information processing facilities"},
    "8.15": {"theme": "8", "name": "Logging"},
    "8.16": {"theme": "8", "name": "Monitoring activities"},
    "8.17": {"theme": "8", "name": "Clock synchronization"},
    "8.18": {"theme": "8", "name": "Use of privileged utility programs"},
    "8.19": {"theme": "8", "name": "Installation of software on operational systems"},
    "8.20": {"theme": "8", "name": "Networks security"},
    "8.21": {"theme": "8", "name": "Security of network services"},
    "8.22": {"theme": "8", "name": "Segregation of networks"},
    "8.23": {"theme": "8", "name": "Web filtering"},
    "8.24": {"theme": "8", "name": "Use of cryptography"},
    "8.25": {"theme": "8", "name": "Secure development life cycle"},
    "8.26": {"theme": "8", "name": "Application security requirements"},
    "8.27": {"theme": "8", "name": "Secure system architecture and engineering principles"},
    "8.28": {"theme": "8", "name": "Secure coding"},
    "8.29": {"theme": "8", "name": "Security testing in development and acceptance"},
    "8.30": {"theme": "8", "name": "Outsourced development"},
    "8.31": {"theme": "8", "name": "Separation of development, test and production environments"},
    "8.32": {"theme": "8", "name": "Change management"},
    "8.33": {"theme": "8", "name": "Test information"},
    "8.34": {"theme": "8", "name": "Protection of information systems during audit testing"},
}

ANNEX_A_EVIDENCE_FILE = OUTPUTS_DIR / "annex_a_evidence.json"
ANNEX_A_STATUSES      = ["Not Assessed", "Implemented", "Partial", "Planned"]

MODEL_GUIDE = [
    {"Model": "phi4-mini:3.8b-q4_K_M", "Best for": "Document generation (default)",
     "VRAM": "GPU+CPU split", "Speed": "~15 tok/s", "min_ram_gb": 0,
     "Install": "ollama pull phi4-mini:3.8b-q4_K_M"},
    {"Model": "qwen2.5:1.5b-q4_K_M", "Best for": "AI Reviewer (fastest)",
     "VRAM": "Fully on GPU (~1.2 GB)", "Speed": "~30 tok/s", "min_ram_gb": 0,
     "Install": "ollama pull qwen2.5:1.5b"},
    {"Model": "llama3.2:3b-q4_K_M", "Best for": "Document generation (alternative)",
     "VRAM": "GPU+CPU split", "Speed": "~18 tok/s", "min_ram_gb": 8,
     "Install": "ollama pull llama3.2:3b"},
    {"Model": "mistral:7b-q4_K_M", "Best for": "High-end machines (8 GB+ VRAM)",
     "VRAM": "~5 GB VRAM needed", "Speed": "~8 tok/s", "min_ram_gb": 16,
     "Install": "ollama pull mistral:7b"},
]


@st.cache_data(ttl=300)
def detect_hardware() -> dict:
    """Detect RAM, VRAM, and CPU. Returns dict with ram_gb, vram_gb, cpu."""
    try:
        import psutil
        ram_gb = round(psutil.virtual_memory().total / 1_073_741_824, 1)
    except Exception:
        ram_gb = 0
    cpu = platform.processor() or platform.machine() or "Unknown CPU"
    vram_gb = 0
    try:
        out = subprocess.check_output(
            ["nvidia-smi", "--query-gpu=memory.total", "--format=csv,noheader,nounits"],
            timeout=3, stderr=subprocess.DEVNULL,
        )
        vram_gb = round(int(out.decode().strip().split()[0]) / 1024, 1)
    except Exception:
        pass
    return {"ram_gb": ram_gb, "vram_gb": vram_gb, "cpu": cpu}

# ---------------------------------------------------------------------------
# Data helpers
# ---------------------------------------------------------------------------

def load_config():
    import os
    with open(CONFIG_PATH, encoding="utf-8") as f:
        cfg = yaml.safe_load(f)
    if url := os.environ.get("OLLAMA_BASE_URL"):
        cfg["llm"]["base_url"] = url
    return cfg

def save_config(cfg):
    with open(CONFIG_PATH, "w", encoding="utf-8") as f:
        yaml.dump(cfg, f, default_flow_style=False, allow_unicode=True)

def load_org():
    if not ORG_PATH.exists():
        return {
            "name": "", "industry": "", "size": "", "scope": "",
            "primary_processes": [], "assets": [], "legal_basis": [],
            "stakeholders": [], "locations": [], "key_personnel": [],
            "critical_suppliers": [], "suppliers": [],
            "existing_controls": [], "certifications_existing": [],
        }
    with open(ORG_PATH, encoding="utf-8") as f:
        return json.load(f)

def save_org(data):
    with open(ORG_PATH, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

def get_clause_status(cid):
    if not (OUTPUTS_DIR / f"{cid}.md").exists():
        return "MISSING"
    sf = OUTPUTS_DIR / f"{cid}.status.json"
    if sf.exists():
        try:
            return json.loads(sf.read_text(encoding="utf-8")).get("status", "DRAFT").upper()
        except Exception:
            pass
    return "DRAFT"

def save_status(cid, status, notes=""):
    OUTPUTS_DIR.mkdir(exist_ok=True)
    (OUTPUTS_DIR / f"{cid}.status.json").write_text(
        json.dumps({"clause_id": cid, "status": status, "notes": notes,
                    "timestamp": datetime.now().isoformat()}, indent=2),
        encoding="utf-8",
    )

def get_review_assessment(cid):
    """Returns internal code: PASS / CONDITIONAL PASS / FAIL / UNKNOWN / None"""
    f = OUTPUTS_DIR / f"{cid}.critic.md"
    if not f.exists():
        return None
    for line in f.read_text(encoding="utf-8", errors="replace").splitlines():
        if "**Overall Assessment:**" in line or "**Review Result:**" in line:
            up = line.upper()
            if "CONDITIONAL" in up:  return "CONDITIONAL PASS"
            if "FAIL"        in up:  return "FAIL"
            if "PASS"        in up:  return "PASS"
    return "UNKNOWN"

def get_review_text(cid):
    f = OUTPUTS_DIR / f"{cid}.critic.md"
    return f.read_text(encoding="utf-8", errors="replace") if f.exists() else None

def read_output(cid):
    f = OUTPUTS_DIR / f"{cid}.md"
    return f.read_text(encoding="utf-8", errors="replace") if f.exists() else None

def completion_stats():
    counts = {"APPROVED": 0, "DRAFT": 0, "REVISION": 0, "MISSING": 0}
    for cid in CLAUSE_NAMES:
        counts[get_clause_status(cid)] += 1
    return len(CLAUSE_NAMES), counts

def read_log_tail(n=60):
    if not LOG_FILE.exists():
        return None
    lines = LOG_FILE.read_text(encoding="utf-8", errors="replace").splitlines()
    return "\n".join(lines[-n:])

def get_ollama_models(base_url):
    try:
        r = requests.get(f"{base_url}/api/tags", timeout=4)
        r.raise_for_status()
        return [m["name"] for m in r.json().get("models", [])] or None
    except Exception:
        return None

def run_reviewer_subprocess(cid):
    cmd = [sys.executable, str(BASE_DIR / "critic.py"), "--clause", cid, "--force"]
    try:
        res = subprocess.run(cmd, capture_output=True, text=True,
                             encoding="utf-8", errors="replace",
                             cwd=str(BASE_DIR), timeout=700)
        return res.returncode == 0, res.stdout + res.stderr
    except Exception as e:
        return False, str(e)

# ---------------------------------------------------------------------------
# Embedding model — cached so it loads only once per session
# ---------------------------------------------------------------------------

@st.cache_resource(show_spinner="Loading ISO knowledge base — please wait...")
def get_embedding_model():
    from sentence_transformers import SentenceTransformer
    return SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")

# ---------------------------------------------------------------------------
# Export helpers
# ---------------------------------------------------------------------------

def _get_personnel_for_doc(org):
    """Return (prepared_by, reviewed_by, approved_by) strings from org key_personnel."""
    personnel = org.get("key_personnel", [])
    if not personnel:
        return "", "", ""

    def _find(keywords):
        for p in personnel:
            role = p.get("role", "").lower()
            if any(kw in role for kw in keywords):
                return f"{p.get('name','')} ({p.get('role','')})"
        return ""

    prepared  = _find(["isms", "information security", "ciso", "security manager", "iso lead"])
    reviewed  = _find(["it manager", "risk owner", "risk manager", "it director"])
    approved  = _find(["ceo", "managing director", "general manager", "top management", "director", "president"])
    if not prepared and len(personnel) > 0:
        prepared = f"{personnel[0].get('name','')} ({personnel[0].get('role','')})"
    if not reviewed and len(personnel) > 1:
        reviewed = f"{personnel[1].get('name','')} ({personnel[1].get('role','')})"
    if not approved and len(personnel) > 2:
        approved = f"{personnel[2].get('name','')} ({personnel[2].get('role','')})"
    return prepared, reviewed, approved

_INLINE_RE = re.compile(r"(\*\*[^*\n]+\*\*|_[^_\n]+_|\*[^*\n]+\*|`[^`\n]+`)")

def _add_inline(paragraph, text):
    from docx.shared import Pt
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif (part.startswith("_") and part.endswith("_")) or \
             (part.startswith("*") and part.endswith("*")):
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)

_TABLE_ROW_RE = re.compile(r"^\s*\|")
_TABLE_SEP_RE = re.compile(r"^\s*\|[-:| ]+\|\s*$")

def _extract_tables(content):
    """Parse markdown table blocks from content.
    Returns (cleaned_content, tables) where tables is a list of
    {"title": str, "headers": list[str], "rows": list[list[str]]}.
    Table blocks in cleaned_content are replaced by a placeholder line.
    """
    lines = content.splitlines()
    result_lines = []
    tables = []
    i = 0
    last_heading = None

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("#"):
            last_heading = stripped.lstrip("#").strip()
            result_lines.append(line)
            i += 1
            continue

        if _TABLE_ROW_RE.match(line) and i + 1 < len(lines) and _TABLE_SEP_RE.match(lines[i + 1]):
            block = []
            while i < len(lines) and _TABLE_ROW_RE.match(lines[i]):
                block.append(lines[i].strip())
                i += 1

            headers, rows = [], []
            for bl in block:
                if _TABLE_SEP_RE.match(bl):
                    continue
                cells = [c.strip() for c in bl.strip("|").split("|")]
                if not headers:
                    headers = cells
                else:
                    rows.append(cells)

            if headers:
                title = last_heading or f"Table {len(tables) + 1}"
                tables.append({"title": title, "headers": headers, "rows": rows})
                result_lines.append(f"_[Table: {title} — included in Appendix below]_")
        else:
            result_lines.append(line)
            i += 1

    return "\n".join(result_lines), tables

def _md_line_to_docx(doc, line):
    s = line.rstrip()
    if s.startswith("#### "): doc.add_heading(s[5:], 4)
    elif s.startswith("### "): doc.add_heading(s[4:], 3)
    elif s.startswith("## "):  doc.add_heading(s[3:], 2)
    elif s.startswith("# "):   doc.add_heading(s[2:], 1)
    elif s.startswith(("- ", "* ")):
        _add_inline(doc.add_paragraph(style="List Bullet"), s[2:])
    elif re.match(r"^\d+\.\s", s):
        _add_inline(doc.add_paragraph(style="List Number"), re.sub(r"^\d+\.\s", "", s))
    elif s.startswith("> "):
        p = doc.add_paragraph()
        from docx.shared import Inches
        p.paragraph_format.left_indent = Inches(0.4)
        run = p.add_run(s[2:])
        run.italic = True
    elif s in ("", "---"):
        doc.add_paragraph("")
    elif _TABLE_ROW_RE.match(s) or _TABLE_SEP_RE.match(s):
        pass
    elif re.match(r"^```", s):
        pass  # skip code fence markers — content already stripped upstream
    else:
        _add_inline(doc.add_paragraph(), s)

@st.cache_data(show_spinner=False)
def export_clause_to_word(cid, content, title_override=None,
                          org_name="", prepared_by="", reviewed_by="", approved_by=""):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    today = datetime.now().strftime("%d %B %Y")
    title = title_override or f"Clause {cid} — {CLAUSE_NAMES.get(cid, cid)}"

    section = doc.sections[0]
    if org_name:
        hdr = section.header.paragraphs[0]
        hdr.text = f"{org_name}  ·  ISMS Documentation  ·  CONFIDENTIAL"
        hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hdr.runs[0].font.size = Pt(8)
        hdr.runs[0].font.color.rgb = RGBColor(0x5A, 0x70, 0x90)
    ftr = section.footer.paragraphs[0]
    ftr.text = f"ISO 27001:2022  ·  Generated by VaultISO27 v{__version__}  ·  {today}"
    ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ftr.runs[0].font.size = Pt(8)
    ftr.runs[0].font.color.rgb = RGBColor(0x5A, 0x70, 0x90)

    doc.add_heading(title, 0)

    if org_name or prepared_by or approved_by:
        info_tbl = doc.add_table(rows=5, cols=2)
        info_tbl.style = "Table Grid"
        _info_rows = [
            ("Organization",   org_name or "—"),
            ("Prepared by",    prepared_by or "—"),
            ("Reviewed by",    reviewed_by or "—"),
            ("Approved by",    approved_by or "—"),
            ("Classification",  "CONFIDENTIAL"),
        ]
        for ri, (lbl, val) in enumerate(_info_rows):
            info_tbl.rows[ri].cells[0].text = lbl
            info_tbl.rows[ri].cells[1].text = val
            info_tbl.rows[ri].cells[0].paragraphs[0].runs[0].bold = True
        doc.add_paragraph("")

    # Strip wrapping code fences the LLM may emit (```markdown ... ```)
    _fence_stripped = re.sub(r"^```[a-zA-Z]*\s*$", "", content, flags=re.MULTILINE)
    cleaned_content, tables = _extract_tables(_fence_stripped)
    for line in cleaned_content.splitlines():
        _md_line_to_docx(doc, line)

    if tables:
        doc.add_page_break()
        doc.add_heading("Appendix — Data Tables", 2)
        for tbl_data in tables:
            doc.add_heading(tbl_data["title"], 3)
            n_cols = len(tbl_data["headers"])
            word_tbl = doc.add_table(rows=1 + len(tbl_data["rows"]), cols=n_cols)
            word_tbl.style = "Table Grid"
            hdr_row = word_tbl.rows[0]
            for ci, h in enumerate(tbl_data["headers"]):
                cell = hdr_row.cells[ci]
                cell.text = h
                cell.paragraphs[0].runs[0].bold = True
            for ri, row_vals in enumerate(tbl_data["rows"], 1):
                for ci, val in enumerate(row_vals):
                    if ci < n_cols:
                        word_tbl.rows[ri].cells[ci].text = val
            doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

@st.cache_data(ttl=15, show_spinner=False)
def export_all_to_excel(org_name=""):
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    wb   = openpyxl.Workbook()
    ws   = wb.active
    ws.title = "ISMS Tracker"
    navy     = "0A2342"

    if org_name:
        ws.merge_cells("A1:F1")
        banner = ws.cell(row=1, column=1, value=f"{org_name}  ·  ISO 27001:2022 ISMS Document Tracker")
        banner.font      = Font(bold=True, size=12, color="FFFFFF")
        banner.fill      = PatternFill("solid", fgColor=navy)
        banner.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[1].height = 22
        header_row = 2
    else:
        header_row = 1

    headers = ["Section", "Document Name", "Status", "AI Review", "Words", "Last Modified"]
    for c, h in enumerate(headers, 1):
        cell = ws.cell(row=header_row, column=c, value=h)
        cell.font      = Font(bold=True, color="FFFFFF")
        cell.fill      = PatternFill("solid", fgColor=navy)
        cell.alignment = Alignment(horizontal="center")
    ws.column_dimensions["A"].width = 8
    ws.column_dimensions["B"].width = 35
    ws.column_dimensions["C"].width = 14
    ws.column_dimensions["D"].width = 22
    ws.column_dimensions["E"].width = 8
    ws.column_dimensions["F"].width = 16
    status_fill = {"APPROVED": "D4EDDA", "DRAFT": "FFF3CD",
                   "REVISION": "F8D7DA", "MISSING": "E2E3E5"}
    for ri, (cid, name) in enumerate(CLAUSE_NAMES.items(), header_row + 1):
        status  = get_clause_status(cid)
        review  = get_review_assessment(cid)
        label,_,_ = REVIEW_RESULT.get(review, ("—","info","neutral")) if review else ("—","info","neutral")
        f       = OUTPUTS_DIR / f"{cid}.md"
        words   = len(f.read_text(encoding="utf-8", errors="replace").split()) if f.exists() else 0
        mod     = datetime.fromtimestamp(f.stat().st_mtime).strftime("%Y-%m-%d") if f.exists() else "—"
        for ci, val in enumerate([cid, name, status, label, words, mod], 1):
            cell = ws.cell(row=ri, column=ci, value=val)
            cell.alignment = Alignment(horizontal="left" if ci == 2 else "center")
            if ci == 3:
                cell.fill = PatternFill("solid", fgColor=status_fill.get(status, "FFFFFF"))
        if f.exists():
            raw = f.read_text(encoding="utf-8", errors="replace")
            ws2 = wb.create_sheet(title=cid)
            ws2.cell(row=1, column=1, value=f"{cid}: {name}").font = Font(bold=True, size=13)
            _, tables = _extract_tables(raw)
            for li, line in enumerate(raw.splitlines(), 3):
                ws2.cell(row=li, column=1, value=line)
            ws2.column_dimensions["A"].width = 120
            for tbl_data in tables:
                sheet_name = f"{cid} {tbl_data['title']}"[:31]
                wst = wb.create_sheet(title=sheet_name)
                wst.cell(row=1, column=1, value=f"{cid}: {tbl_data['title']}").font = Font(bold=True, size=12)
                for ci, h in enumerate(tbl_data["headers"], 1):
                    cell = wst.cell(row=3, column=ci, value=h)
                    cell.font = Font(bold=True, color="FFFFFF")
                    cell.fill = PatternFill("solid", fgColor=navy)
                    cell.alignment = Alignment(horizontal="center")
                for ri, row_vals in enumerate(tbl_data["rows"], 4):
                    for ci, val in enumerate(row_vals, 1):
                        if ci <= len(tbl_data["headers"]):
                            wst.cell(row=ri, column=ci, value=val).alignment = Alignment(horizontal="left")
                for col_idx in range(1, len(tbl_data["headers"]) + 1):
                    from openpyxl.utils import get_column_letter
                    wst.column_dimensions[get_column_letter(col_idx)].width = 25
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()

# ---------------------------------------------------------------------------
# Asset register + supplier register parsers
# ---------------------------------------------------------------------------

def _parse_asset_register(file_bytes: bytes) -> list:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)

    ws = None
    for name in wb.sheetnames:
        if any(k in name.lower() for k in ("asset", "list", "register")):
            ws = wb[name]; break
    if ws is None:
        for name in wb.sheetnames:
            if name.lower() not in ("title", "cover", "data"):
                ws = wb[name]; break
    if ws is None:
        return []

    rows = list(ws.iter_rows(values_only=True))

    hdr_idx = None
    for i, row in enumerate(rows):
        cells = [str(c).lower() if c else "" for c in row]
        if any("name" in c or "asset" in c for c in cells):
            hdr_idx = i; break
    if hdr_idx is None:
        return []

    headers = [str(h).strip() if h else "" for h in rows[hdr_idx]]
    col = {}
    for i, h in enumerate(headers):
        hl = h.lower()
        if "asset no" in hl or hl in ("no", "id", "asset_no"):
            col.setdefault("asset_no", i)
        if ("name" in hl and col.get("name") is None) or hl == "name":
            col.setdefault("name", i)
        if "description" in hl or hl == "desc":
            col.setdefault("description", i)
        if "type" in hl:
            col.setdefault("type", i)
        if "responsible" in hl or "owner" in hl:
            col.setdefault("responsible", i)
        if "classification" in hl or "class" in hl:
            col.setdefault("data_classification", i)
        if "risk score" in hl or "risk_score" in hl:
            col.setdefault("risk_score", i)
        if "personal" in hl:
            col.setdefault("personal_data", i)
        if hl == "ai":
            col.setdefault("ai_involved", i)
        if "license" in hl:
            col.setdefault("license_type", i)
        if "user" in hl:
            col.setdefault("users", i)

    assets = []
    for row in rows[hdr_idx + 1:]:
        if not any(c is not None and str(c).strip() for c in row):
            continue
        a = {}
        for key, idx in col.items():
            val = row[idx] if idx < len(row) else None
            if val is not None:
                a[key] = val
        if a.get("name") or a.get("asset_no"):
            assets.append(a)
    return assets


def _parse_supplier_register(file_bytes: bytes) -> list:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)

    ws = None
    for name in wb.sheetnames:
        nl = name.lower()
        if any(k in nl for k in ("risk", "supplier", "evaluation", "vendor")):
            ws = wb[name]; break
    if ws is None:
        for name in wb.sheetnames:
            if name.lower() not in ("title", "cover", "data"):
                ws = wb[name]; break
    if ws is None:
        return []

    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []

    first_col_vals = [str(r[0]).strip().lower() if r[0] else "" for r in rows[:5]]
    is_transposed = any(v in ("field", "responsible") for v in first_col_vals)

    if is_transposed:
        asset_names = [str(rows[0][c]).strip() if rows[0][c] else ""
                       for c in range(1, len(rows[0]))]
        field_rows: dict[str, int] = {}
        for i, row in enumerate(rows[1:], 1):
            key = str(row[0]).strip().lower() if row[0] else ""
            if key:
                field_rows[key] = i

        def _get(col_i: int, *searches) -> str:
            for search in searches:
                for k, ri in field_rows.items():
                    if search in k:
                        v = rows[ri][col_i] if col_i < len(rows[ri]) else None
                        return str(v).strip() if v is not None else ""
            return ""

        suppliers = []
        for ci, asset_name in enumerate(asset_names, 1):
            if not asset_name:
                continue
            involved_raw = _get(ci, "external service")
            involved = involved_raw.lower() in ("yes", "true", "1")
            s = {
                "asset": asset_name,
                "provider_involved": involved,
                "services": _get(ci, "services provided"),
                "provider_activities": _get(ci, "provider activities"),
                "dependency_level": _get(ci, "dependency on"),
                "reason_for_dependency": _get(ci, "reason for"),
                "processed_data": _get(ci, "processed data"),
                "data_classification": _get(ci, "data classification"),
            }
            if involved or s["services"]:
                suppliers.append(s)
        return suppliers

    hdr_idx = None
    for i, row in enumerate(rows):
        cells = [str(c).lower() if c else "" for c in row]
        if any(k in c for c in cells for k in ("name", "supplier", "vendor")):
            hdr_idx = i; break
    if hdr_idx is None:
        return []

    headers = [str(h).strip() if h else "" for h in rows[hdr_idx]]
    col = {}
    for i, h in enumerate(headers):
        hl = h.lower()
        if any(k in hl for k in ("name", "supplier", "vendor")):
            col.setdefault("asset", i)
        if "service" in hl:
            col.setdefault("services", i)
        if "classification" in hl:
            col.setdefault("data_classification", i)
        if "depend" in hl:
            col.setdefault("dependency_level", i)

    suppliers = []
    for row in rows[hdr_idx + 1:]:
        if not any(c is not None and str(c).strip() for c in row):
            continue
        s = {k: (str(row[v]).strip() if v < len(row) and row[v] is not None else "")
             for k, v in col.items()}
        if any(s.values()):
            s["provider_involved"] = True
            suppliers.append(s)
    return suppliers


# ---------------------------------------------------------------------------
# Organization extraction
# ---------------------------------------------------------------------------

def extract_text_from_upload(uploaded_file):
    suffix = Path(uploaded_file.name).suffix.lower()
    if suffix == ".txt":
        return uploaded_file.read().decode("utf-8", errors="replace")
    elif suffix == ".pdf":
        try:
            import pypdf
            reader = pypdf.PdfReader(uploaded_file)
            return "\n".join(p.extract_text() or "" for p in reader.pages)
        except Exception as e:
            st.error(f"Could not read PDF: {e}")
            return None
    elif suffix in (".docx", ".doc"):
        try:
            from docx import Document as DocxDoc
            doc = DocxDoc(uploaded_file)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            st.error(f"Could not read Word document: {e}")
            return None
    st.error(f"File type not supported: {suffix}. Please upload a PDF, Word (.docx), or text (.txt) file.")
    return None


ORG_JSON_SCHEMA = {
    "name": "", "industry": "", "size": "", "scope": "",
    "primary_processes": [], "locations": [], "legal_basis": [],
    "stakeholders": [{"name": "", "expectation": ""}],
    "assets": [{"name": "", "system": "", "owner": "", "classification": ""}],
    "key_personnel": [{"role": "", "name": ""}],
    "critical_suppliers": [], "existing_controls": [], "certifications_existing": [],
}


def extract_org_with_llm(text, cfg):
    schema_str = json.dumps(ORG_JSON_SCHEMA, indent=2)
    prompt = f"""You are an ISO 27001 consultant. Extract organization information from the document below.

Return ONLY a valid JSON object with this exact structure — no explanation, no markdown fences:

{schema_str}

Rules:
- "scope": 1–2 sentences describing the core business and IT activities suitable for an ISMS scope statement
- "size": format as "45 employees" or "~100 employees" if approximate
- "legal_basis": include only regulations explicitly mentioned in the document (e.g. GDPR, NIS2, SOC2, PCI-DSS, ISO 27001) — do NOT invent or assume any regulation not stated in the text
- "existing_controls": list specific security controls mentioned (MFA, VPN, firewalls, encryption, etc.)
- Leave fields as empty string or empty array when not found in the document
- Do NOT invent information

DOCUMENT:
{text[:5000]}

JSON:"""

    try:
        requests.get(f"{cfg['llm']['base_url']}/api/tags", timeout=5)
    except Exception:
        st.error("Cannot reach the AI engine. "
                 "Go to Organization > AI Engine and check that Ollama is running.")
        return None

    try:
        resp = requests.post(
            f"{cfg['llm']['base_url']}/api/generate",
            json={"model": cfg["llm"]["model"], "prompt": prompt, "stream": False,
                  "options": {"temperature": 0.05, "num_predict": 600}},
            timeout=180,
        )
        raw = resp.json().get("response", "").strip()
        start, end = raw.find("{"), raw.rfind("}") + 1
        if start >= 0 and end > start:
            return json.loads(raw[start:end])
        st.error("The AI returned an unexpected response. "
                 "Try uploading a shorter or plain-text (.txt) version of your document.")
    except requests.exceptions.Timeout:
        st.error(
            "The AI engine is taking too long (3 minute limit reached). "
            "This usually means the model is still loading into memory. "
            "Wait 1–2 minutes, then try again — it will be faster on the second attempt."
        )
    except requests.exceptions.ConnectionError:
        st.error("Cannot reach the AI engine. "
                 "Go to Organization > AI Engine and check that Ollama is running.")
    except json.JSONDecodeError:
        st.error("Unexpected AI response format. "
                 "Try uploading a shorter or simpler document (plain text works best).")
    except Exception as e:
        st.error(f"Extraction error: {e}")
    return None


def extract_personnel_with_llm(text, cfg):
    """Extract key personnel names and roles from an org chart or document."""
    prompt = f"""You are an ISO 27001 consultant. Extract key personnel information from the document below.

Return ONLY a valid JSON array — no explanation, no markdown fences:
[
  {{"role": "CEO", "name": "Full Name"}},
  {{"role": "CISO", "name": "Full Name"}}
]

Rules:
- Include only named individuals with clear roles
- Roles should map to information security governance (CEO, CISO, IT Manager, Risk Owner, DPO, etc.)
- Return an empty array [] if no clear personnel found
- Do NOT invent names

DOCUMENT:
{text[:3000]}

JSON:"""
    try:
        resp = requests.post(
            f"{cfg['llm']['base_url']}/api/generate",
            json={"model": cfg["llm"]["model"], "prompt": prompt, "stream": False,
                  "options": {"temperature": 0.05, "num_predict": 300}},
            timeout=120,
        )
        raw = resp.json().get("response", "").strip()
        start, end = raw.find("["), raw.rfind("]") + 1
        if start >= 0 and end > start:
            return json.loads(raw[start:end])
    except Exception as e:
        st.error(f"Personnel extraction error: {e}")
    return None


# ---------------------------------------------------------------------------
# GitHub upload
# ---------------------------------------------------------------------------

def upload_clause_to_github(cid, content, cfg):
    """Upload a single clause markdown file to GitHub. Returns (ok, message)."""
    gh_cfg = cfg.get("github", {})
    token  = gh_cfg.get("token", "")
    repo   = gh_cfg.get("repo", "")
    folder = gh_cfg.get("folder", "isms-docs/").rstrip("/")

    if not token or not repo:
        return False, "GitHub token or repository not configured in Settings."
    try:
        from github import Github, GithubException
    except ImportError:
        return False, "PyGithub not installed. Run: pip install PyGithub"

    try:
        g        = Github(token)
        gh_repo  = g.get_repo(repo)
        path     = f"{folder}/{cid}.md"
        message  = f"Update ISMS clause {cid} — {CLAUSE_NAMES.get(cid, cid)}"
        encoded  = content.encode("utf-8")
        try:
            existing = gh_repo.get_contents(path)
            gh_repo.update_file(path, message, encoded, existing.sha)
            return True, f"Updated `{path}` in `{repo}`."
        except GithubException:
            gh_repo.create_file(path, message, encoded)
            return True, f"Created `{path}` in `{repo}`."
    except Exception as e:
        return False, str(e)


# ---------------------------------------------------------------------------
# Utility
# ---------------------------------------------------------------------------

def _field_val(v):
    """Return True if a field value is non-empty."""
    if isinstance(v, list):
        return any(
            (isinstance(i, str) and i.strip()) or
            (isinstance(i, dict) and any(str(x).strip() for x in i.values()))
            for i in v
        )
    return bool(str(v).strip())


# ---------------------------------------------------------------------------
# Annex A helpers
# ---------------------------------------------------------------------------

def load_annex_a() -> dict:
    if ANNEX_A_EVIDENCE_FILE.exists():
        try:
            return json.loads(ANNEX_A_EVIDENCE_FILE.read_text(encoding="utf-8"))
        except Exception:
            pass
    return {}

def save_annex_a(data: dict) -> None:
    OUTPUTS_DIR.mkdir(exist_ok=True)
    ANNEX_A_EVIDENCE_FILE.write_text(
        json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8"
    )

@st.cache_data(show_spinner=False)
def export_soa_to_excel(evidence_json: str, org_name: str = "") -> bytes:
    """Generate Statement of Applicability Excel. evidence_json is the raw JSON string for caching."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    evidence = json.loads(evidence_json) if evidence_json else {}
    wb  = openpyxl.Workbook()
    ws  = wb.active
    ws.title = "Statement of Applicability"

    navy    = "0A2342"
    thin    = Side(style="thin", color="CCCCCC")
    border  = Border(left=thin, right=thin, top=thin, bottom=thin)

    STATUS_FILL = {
        "Implemented":  "C6EFCE",
        "Partial":      "FFEB9C",
        "Planned":      "FCE4D6",
        "Not Assessed": "F2F2F2",
    }

    row = 1
    if org_name:
        ws.merge_cells(f"A{row}:G{row}")
        c = ws.cell(row=row, column=1,
                    value=f"{org_name}  ·  ISO 27001:2022 Statement of Applicability")
        c.font      = Font(bold=True, size=12, color="FFFFFF")
        c.fill      = PatternFill("solid", fgColor=navy)
        c.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[row].height = 22
        row += 1

    headers = [
        "Control ID", "Control Name", "Theme", "Applicable",
        "Status", "Justification / Notes", "Evidence References",
    ]
    col_widths = [12, 55, 22, 12, 16, 40, 40]
    for ci, (h, w) in enumerate(zip(headers, col_widths), 1):
        c = ws.cell(row=row, column=ci, value=h)
        c.font      = Font(bold=True, color="FFFFFF")
        c.fill      = PatternFill("solid", fgColor=navy)
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        c.border    = border
        ws.column_dimensions[get_column_letter(ci)].width = w
    ws.row_dimensions[row].height = 20
    header_row = row
    row += 1

    ws.freeze_panes = f"A{header_row + 1}"

    for cid, ctrl in ANNEX_A_CONTROLS.items():
        e          = evidence.get(cid, {})
        applicable = e.get("applicable", True)
        status_raw = e.get("status", "Not Assessed") if applicable else "N/A"
        justif     = e.get("justification", "")
        evid_refs  = "\n".join(e.get("evidence_refs", []))
        theme_name = ANNEX_A_THEMES.get(ctrl["theme"], ctrl["theme"])

        fill_key  = status_raw if status_raw in STATUS_FILL else "Not Assessed"
        row_fill  = PatternFill("solid", fgColor=STATUS_FILL[fill_key]) if applicable else \
                    PatternFill("solid", fgColor="D9D9D9")

        STATUS_TRANSLATE = {
            "Implemented": "Implemented", "Partial": "Partial",
            "Planned": "Planned", "Not Assessed": "Not Assessed",
        }
        status = STATUS_TRANSLATE.get(status_raw, status_raw)

        vals = [cid, ctrl["name"], theme_name,
                "Yes" if applicable else "No",
                status, justif, evid_refs]
        for ci, val in enumerate(vals, 1):
            c = ws.cell(row=row, column=ci, value=val)
            c.fill      = row_fill
            c.border    = border
            c.alignment = Alignment(vertical="top", wrap_text=(ci >= 6))
        ws.row_dimensions[row].height = 18
        row += 1

    ws.auto_filter.ref = f"A{header_row}:G{row - 1}"

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _annex_collect_from_state() -> dict:
    """Collect all annex widget values from st.session_state into a dict."""
    data = {}
    for cid in ANNEX_A_CONTROLS:
        applicable = st.session_state.get(f"annex_{cid}_applicable", True)
        data[cid] = {
            "applicable":    applicable,
            "status":        st.session_state.get(f"annex_{cid}_status", "Not Assessed"),
            "justification": st.session_state.get(f"annex_{cid}_justification", ""),
            "evidence_refs": [
                r.strip() for r in
                st.session_state.get(f"annex_{cid}_evidence", "").splitlines()
                if r.strip()
            ],
        }
    return data
