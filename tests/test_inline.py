"""Quick smoke-test for _add_inline and _extract_tables without running Streamlit."""
import re
import sys
import io
from pathlib import Path
from datetime import datetime

# ── inline renderer ──────────────────────────────────────────────────────────
_INLINE_RE = re.compile(r"(\*\*[^*\n]+\*\*|_[^_\n]+_|\*[^*\n]+\*|`[^`\n]+`)")


def _add_inline(paragraph, text):
    from docx.shared import Pt
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif (part.startswith("_") and part.endswith("_")) or \
             (part.startswith("*") and part.endswith("*")):
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


# ── table extraction ─────────────────────────────────────────────────────────
_TABLE_ROW_RE = re.compile(r"^\s*\|")
_TABLE_SEP_RE = re.compile(r"^\s*\|[-:| ]+\|\s*$")


def _extract_tables(content):
    lines = content.splitlines()
    result_lines = []
    tables = []
    i = 0
    last_heading = None
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("#"):
            last_heading = stripped.lstrip("#").strip()
            result_lines.append(line)
            i += 1
            continue
        if _TABLE_ROW_RE.match(line) and i + 1 < len(lines) and _TABLE_SEP_RE.match(lines[i + 1]):
            block = []
            while i < len(lines) and _TABLE_ROW_RE.match(lines[i]):
                block.append(lines[i].strip())
                i += 1
            headers, rows = [], []
            for bl in block:
                if _TABLE_SEP_RE.match(bl):
                    continue
                cells = [c.strip() for c in bl.strip("|").split("|")]
                if not headers:
                    headers = cells
                else:
                    rows.append(cells)
            if headers:
                title = last_heading or f"Table {len(tables) + 1}"
                tables.append({"title": title, "headers": headers, "rows": rows})
                result_lines.append(f"_[Table: {title} — included in Appendix below]_")
        else:
            result_lines.append(line)
            i += 1
    return "\n".join(result_lines), tables


# ── tests ────────────────────────────────────────────────────────────────────

def test_bold():
    from docx import Document
    doc = Document()
    p = doc.add_paragraph()
    _add_inline(p, "Normal **bold text** after")
    runs = p.runs
    assert runs[1].bold is True
    assert runs[1].text == "bold text"


def test_italic_underscore():
    from docx import Document
    doc = Document()
    p = doc.add_paragraph()
    _add_inline(p, "Before _italic_ after")
    bold_runs = [r for r in p.runs if r.italic]
    assert len(bold_runs) == 1
    assert bold_runs[0].text == "italic"


def test_inline_code():
    from docx import Document
    doc = Document()
    p = doc.add_paragraph()
    _add_inline(p, "Use `python main.py` to run")
    code_runs = [r for r in p.runs if r.font.name == "Courier New"]
    assert len(code_runs) == 1
    assert code_runs[0].text == "python main.py"


def test_extract_tables_basic():
    content = (
        "## Risk Assessment\n\n"
        "### Risk Register\n\n"
        "| Risk ID | Threat | Treatment |\n"
        "|---------|--------|-----------|\n"
        "| R-001   | Phishing | MFA     |\n\n"
        "Trailing text.\n"
    )
    cleaned, tables = _extract_tables(content)
    assert len(tables) == 1
    assert tables[0]["title"] == "Risk Register"
    assert tables[0]["headers"] == ["Risk ID", "Threat", "Treatment"]
    assert tables[0]["rows"][0] == ["R-001", "Phishing", "MFA"]
    assert "| R-001" not in cleaned
    assert "Risk Register" in cleaned


def test_extract_tables_no_table():
    content = "## Section\n\nJust prose here.\n"
    cleaned, tables = _extract_tables(content)
    assert tables == []
    assert cleaned == content.rstrip("\n")


def test_extract_tables_multiple():
    content = (
        "### Asset Register\n\n"
        "| Asset | Owner |\n"
        "|-------|-------|\n"
        "| CRM   | IT    |\n\n"
        "### Risk Register\n\n"
        "| Risk | Level |\n"
        "|------|-------|\n"
        "| SQL  | High  |\n"
    )
    _, tables = _extract_tables(content)
    assert len(tables) == 2
    assert tables[0]["title"] == "Asset Register"
    assert tables[1]["title"] == "Risk Register"


def test_bold_renders_not_stripped():
    """Bold text must produce a bold run, not literal asterisks."""
    from docx import Document
    doc = Document()
    p = doc.add_paragraph()
    _add_inline(p, "The **ISMS scope** includes all systems.")
    bold_runs = [r for r in p.runs if r.bold]
    assert len(bold_runs) == 1
    assert bold_runs[0].text == "ISMS scope"
    full_text = "".join(r.text for r in p.runs)
    assert "**" not in full_text


def test_fence_markers_skipped_by_md_line_to_docx():
    """Code fence markers should produce no paragraph in the Word document."""
    import re
    from docx import Document

    _TABLE_ROW_RE = re.compile(r"^\s*\|")
    _TABLE_SEP_RE = re.compile(r"^\s*\|[-:| ]+\|\s*$")

    def _md_line_to_docx(doc, line):
        s = line.rstrip()
        if s.startswith("#### "):
            doc.add_heading(s[5:], 4)
        elif s.startswith("### "):
            doc.add_heading(s[4:], 3)
        elif s.startswith("## "):
            doc.add_heading(s[3:], 2)
        elif s.startswith("# "):
            doc.add_heading(s[2:], 1)
        elif s.startswith(("- ", "* ")):
            _add_inline(doc.add_paragraph(style="List Bullet"), s[2:])
        elif re.match(r"^\d+\.\s", s):
            _add_inline(doc.add_paragraph(style="List Number"), re.sub(r"^\d+\.\s", "", s))
        elif s.startswith("> "):
            p = doc.add_paragraph()
            p.add_run(s[2:]).italic = True
        elif s in ("", "---"):
            doc.add_paragraph("")
        elif _TABLE_ROW_RE.match(s) or _TABLE_SEP_RE.match(s):
            pass
        elif re.match(r"^```", s):
            pass  # fence markers skipped
        else:
            _add_inline(doc.add_paragraph(), s)

    doc = Document()
    para_count_before = len(doc.paragraphs)
    _md_line_to_docx(doc, "```markdown")
    _md_line_to_docx(doc, "```")
    # Fence lines must not add any paragraphs
    assert len(doc.paragraphs) == para_count_before


def test_fence_stripped_before_export():
    """Pre-pass in export should remove ``` fences from content."""
    import re
    content = "```markdown\n# Scope\n\nBody text here.\n```"
    stripped = re.sub(r"^```[a-zA-Z]*\s*$", "", content, flags=re.MULTILINE)
    assert "```" not in stripped
    assert "# Scope" in stripped
    assert "Body text here." in stripped
